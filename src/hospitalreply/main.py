import os
import re
import copy
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess

pattern = r"尊敬的(.+?)："

patternMatch = r"xxx"

source_dir = os.path.abspath("src/public")
target_dir = os.path.abspath("src/dist")

links = []


def updateDoc(docPath):
    doc = Document(docPath)
    templateDoc = Document("C:/Users/hongxian.yi/Desktop/template.docx")

    found_name = None
    for para in doc.paragraphs:
        match = re.search(pattern, para.text)
        if match:
            found_name = match.group(1)
            break

    if not found_name:
        print("没有找到用户名", filePath)
        return

    found_image = False
    have_text = False

    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

    # 添加分页符和头子
    for para in doc.paragraphs:
        if have_text:
            break
        if found_image:
            for templatePara in templateDoc.paragraphs:
                for templateRun in templatePara.runs:
                    if re.search(patternMatch, templateRun.text):
                        templateRun.text = templateRun.text.replace("xxx", found_name)
                para._element.addprevious(copy.deepcopy(templatePara._element))
            break
        for run in para.runs:
            if run.text:
                have_text = True
                break
            if "graphic" in run._element.xml:  # 检查 run 中是否包含图片
                found_image = True

                # 3. 在图片后添加分页符
                run.add_break(WD_BREAK.PAGE)

                break

    for para in doc.paragraphs:
        have_image = False
        for run in para.runs:
            if "graphic" in run._element.xml:  # 检查 run 中是否包含图片
                have_image = True
                break
        if have_image:
            para.paragraph_format.left_indent = None  # 设置左缩进为0
            para.paragraph_format.right_indent = None  # 设置右缩进为0
            para.paragraph_format.first_line_indent = None  # 设置首行缩进为0
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if have_text:
        print("文字前没有图片", filePath)
    else:
        targetPath = filePath.replace(source_dir, target_dir)
        targetPathDir = os.path.dirname(targetPath)
        if not os.path.exists(targetPathDir):
            os.makedirs(targetPathDir)
        doc.save(filePath.replace(source_dir, target_dir))


for root, dirs, files in os.walk(source_dir):
    for file in files:
        filePath = os.path.join(root, file)
        if not filePath.endswith(".docx"):
            sp = subprocess.run(
                [
                    "C:\Program Files\LibreOffice\program\soffice.exe",
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    os.path.dirname(filePath),
                    filePath,
                ]
            )
            if sp.returncode == 0:
                os.remove(filePath)
                print("转换成功，已删除", filePath)
            else:
                print("命令执行失败:", filePath)
                print("错误信息:", sp.stderr)  # 打印命令的标准错误

count = 1
for root, dirs, files in os.walk(source_dir):
    for file in files:
        filePath = os.path.join(root, file)
        if filePath.endswith(".docx"):
            count = count + 1
            finalFilePath = filePath.replace(source_dir, target_dir).replace("\\", "/")
            links.append((finalFilePath, "file:///" + finalFilePath))

            # try:
            updateDoc(filePath)
            # except:
            #     print("!!!",filePath);

print("总共：", count)

with open("./file.md", "w", encoding="utf-8") as f:
    # 遍历每个链接，按 Markdown 格式写入
    for link_text, url in links:
        f.write(f"[{link_text}]({url})\n")
